"""Export book projects to Markdown, plain text, DOCX, and print-ready PDF."""

from __future__ import annotations

from html import escape
from pathlib import Path

from docx import Document
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)

from ai_book_batch_writer.models import BookProject
from ai_book_batch_writer.project_store import save_project
from ai_book_batch_writer.utils import ensure_parent_dir


def _title(project: BookProject) -> str:
    return project.settings.title or "Untitled Book"


def _metadata_lines(project: BookProject) -> list[tuple[str, str]]:
    settings = project.settings
    return [
        ("Language", settings.output_language),
        ("Tone", settings.tone),
        ("Target audience", settings.target_audience),
        ("Main idea", settings.idea),
        ("Outline tokens", str(project.token_usage.outline_tokens)),
        ("Book tokens", str(project.token_usage.book_tokens)),
        ("Total tokens", str(project.token_usage.total_tokens)),
    ]


def render_markdown(project: BookProject) -> str:
    """Render a project as Markdown."""
    lines = [f"# {_title(project)}"]
    if project.subtitle:
        lines.extend(["", f"*{project.subtitle}*"])

    lines.extend(["", "## Project Metadata", ""])
    lines.extend(f"- **{label}:** {value}" for label, value in _metadata_lines(project))
    lines.extend(["", "## Table of Contents", ""])
    lines.extend(
        f"{chapter.number}. [{chapter.title}](#chapter-{chapter.number})"
        for chapter in project.chapters
    )

    for chapter in project.chapters:
        lines.extend(
            [
                "",
                f'<a id="chapter-{chapter.number}"></a>',
                f"## Chapter {chapter.number}: {chapter.title}",
                "",
            ]
        )
        if chapter.content:
            lines.extend([chapter.content.strip(), ""])
        for section in chapter.sections:
            lines.extend([f"### {section.title}", ""])
            if section.content:
                lines.extend([section.content.strip(), ""])
            else:
                lines.extend([f"*{section.summary}*", ""])
    return "\n".join(lines).rstrip() + "\n"


def export_markdown(project: BookProject, path: str | Path) -> None:
    """Export a project to a UTF-8 Markdown file."""
    target = ensure_parent_dir(path)
    target.write_text(render_markdown(project), encoding="utf-8")


def export_txt(project: BookProject, path: str | Path) -> None:
    """Export a project to a readable UTF-8 plain-text file."""
    lines = [_title(project)]
    if project.subtitle:
        lines.append(project.subtitle)
    lines.append("=" * max(12, len(lines[0])))
    lines.append("")
    lines.extend(f"{label}: {value}" for label, value in _metadata_lines(project))
    lines.extend(["", "TABLE OF CONTENTS", ""])
    lines.extend(
        f"{chapter.number}. {chapter.title}" for chapter in project.chapters
    )

    for chapter in project.chapters:
        heading = f"CHAPTER {chapter.number}: {chapter.title}"
        lines.extend(["", heading, "-" * len(heading), ""])
        if chapter.content:
            lines.extend([chapter.content.strip(), ""])
        for section in chapter.sections:
            lines.extend([section.title, "~" * len(section.title), ""])
            lines.extend([(section.content or section.summary).strip(), ""])

    target = ensure_parent_dir(path)
    target.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def export_docx(project: BookProject, path: str | Path) -> None:
    """Export a project to DOCX with semantic heading levels."""
    document = Document()
    document.add_heading(_title(project), level=1)
    if project.subtitle:
        document.add_paragraph(project.subtitle, style="Subtitle")

    document.add_heading("Project Metadata", level=2)
    for label, value in _metadata_lines(project):
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{label}: ").bold = True
        paragraph.add_run(value)

    document.add_heading("Table of Contents", level=2)
    for chapter in project.chapters:
        document.add_paragraph(
            f"{chapter.number}. {chapter.title}",
            style="List Number",
        )

    for chapter in project.chapters:
        document.add_heading(
            f"Chapter {chapter.number}: {chapter.title}",
            level=2,
        )
        if chapter.content:
            for paragraph_text in chapter.content.split("\n\n"):
                if paragraph_text.strip():
                    document.add_paragraph(paragraph_text.strip())
        for section in chapter.sections:
            document.add_heading(section.title, level=3)
            content = section.content or section.summary
            for paragraph_text in content.split("\n\n"):
                if paragraph_text.strip():
                    document.add_paragraph(paragraph_text.strip())

    target = ensure_parent_dir(path)
    document.save(target)


def export_pdf(project: BookProject, path: str | Path) -> None:
    """Export a clean reading/printing PDF without project metadata."""
    target = ensure_parent_dir(path)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "BookTitle",
        parent=styles["Title"],
        alignment=TA_CENTER,
        fontSize=25,
        leading=31,
        spaceAfter=12,
    )
    subtitle_style = ParagraphStyle(
        "BookSubtitle",
        parent=styles["BodyText"],
        alignment=TA_CENTER,
        fontSize=13,
        leading=18,
        textColor="#555555",
    )
    chapter_style = ParagraphStyle(
        "BookChapter",
        parent=styles["Heading1"],
        fontSize=19,
        leading=24,
        spaceAfter=12,
    )
    section_style = ParagraphStyle(
        "BookSection",
        parent=styles["Heading2"],
        fontSize=14,
        leading=19,
        spaceBefore=12,
        spaceAfter=8,
    )
    body_style = ParagraphStyle(
        "BookBody",
        parent=styles["BodyText"],
        fontSize=10.5,
        leading=15,
        spaceAfter=8,
    )

    story: list[object] = [
        Spacer(1, 55 * mm),
        Paragraph(escape(_title(project)), title_style),
    ]
    if project.subtitle:
        story.append(Paragraph(escape(project.subtitle), subtitle_style))
    story.extend([PageBreak(), Paragraph("Contents", chapter_style)])
    for chapter in project.chapters:
        story.append(
            Paragraph(
                f"{chapter.number}. {escape(chapter.title)}",
                body_style,
            )
        )
    story.append(PageBreak())

    for index, chapter in enumerate(project.chapters):
        if index:
            story.append(PageBreak())
        story.append(
            Paragraph(
                f"Chapter {chapter.number}: {escape(chapter.title)}",
                chapter_style,
            )
        )
        if chapter.content:
            story.extend(_pdf_paragraphs(chapter.content, body_style))
        for section in chapter.sections:
            story.append(Paragraph(escape(section.title), section_style))
            story.extend(
                _pdf_paragraphs(
                    section.content or section.summary,
                    body_style,
                )
            )

    document = SimpleDocTemplate(
        str(target),
        pagesize=A4,
        rightMargin=22 * mm,
        leftMargin=22 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
        title=_title(project),
        author="",
        subject="",
    )
    document.build(
        story,
        onFirstPage=_draw_page_number,
        onLaterPages=_draw_page_number,
    )


def export_project_bundle(
    project: BookProject,
    directory: str | Path,
) -> dict[str, Path]:
    """Save a complete project and all supported exports in one directory."""
    folder = Path(directory)
    folder.mkdir(parents=True, exist_ok=True)
    paths = {
        "json": folder / "project.json",
        "md": folder / "book.md",
        "txt": folder / "book.txt",
        "docx": folder / "book.docx",
        "pdf": folder / "book.pdf",
    }
    save_project(project, paths["json"])
    export_markdown(project, paths["md"])
    export_txt(project, paths["txt"])
    export_docx(project, paths["docx"])
    export_pdf(project, paths["pdf"])
    return paths


def _pdf_paragraphs(
    content: str,
    style: ParagraphStyle,
) -> list[Paragraph]:
    paragraphs = []
    for text in content.split("\n\n"):
        cleaned = text.strip()
        if cleaned:
            paragraphs.append(
                Paragraph(escape(cleaned).replace("\n", "<br/>"), style)
            )
    return paragraphs


def _draw_page_number(canvas, document) -> None:
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor("#666666")
    canvas.drawCentredString(A4[0] / 2, 10 * mm, str(document.page))
    canvas.restoreState()
